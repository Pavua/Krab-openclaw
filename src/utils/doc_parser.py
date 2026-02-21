# -*- coding: utf-8 -*-
"""
Document Intelligence Engine.
Парсинг и анализ документов: PDF, DOCX, XLSX, TXT, CSV.

Зачем: Позволяет боту читать и анализировать документы, отправленные
в Telegram, извлекать текст и индексировать его в RAG.
Связь: Вызывается из main.py при получении документов,
результат индексируется в rag_engine.py.
"""

import os
import logging
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger("DocIntelligence")


class DocumentParser:
    """
    Универсальный парсер документов.
    Поддерживает: PDF, DOCX, XLSX, CSV, TXT, JSON, Markdown.
    """
    
    # Максимальный размер извлекаемого текста (чтобы не перегрузить промпт)
    MAX_TEXT_LENGTH = 15000
    
    # Поддерживаемые MIME-типы и расширения
    SUPPORTED_EXTENSIONS = {
        '.pdf', '.docx', '.doc', '.xlsx', '.xls', 
        '.csv', '.txt', '.md', '.json', '.py', '.js',
        '.html', '.xml', '.yaml', '.yml', '.toml', '.ini',
        '.log', '.conf', '.cfg', '.env'
    }
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Проверяет, поддерживается ли формат файла."""
        ext = Path(filename).suffix.lower()
        return ext in cls.SUPPORTED_EXTENSIONS
    
    @classmethod
    async def parse(cls, file_path: str) -> Tuple[str, dict]:
        """
        Парсит документ и возвращает (текст, метаданные).
        
        Returns:
            tuple: (extracted_text, metadata_dict)
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        metadata = {
            "filename": path.name,
            "extension": ext,
            "size_kb": round(path.stat().st_size / 1024, 1),
            "source": "document_parser"
        }
        
        try:
            if ext == '.pdf':
                text = await cls._parse_pdf(file_path)
            elif ext in ('.docx', '.doc'):
                text = await cls._parse_docx(file_path)
            elif ext in ('.xlsx', '.xls'):
                text = await cls._parse_xlsx(file_path)
            elif ext == '.csv':
                text = await cls._parse_csv(file_path)
            elif ext in ('.txt', '.md', '.json', '.py', '.js', '.html', 
                         '.xml', '.yaml', '.yml', '.toml', '.ini',
                         '.log', '.conf', '.cfg', '.env'):
                text = await cls._parse_text(file_path)
            else:
                return f"⚠️ Формат {ext} пока не поддерживается.", metadata
            
            # Обрезаем если слишком длинный
            if len(text) > cls.MAX_TEXT_LENGTH:
                text = text[:cls.MAX_TEXT_LENGTH] + "\n\n... [текст обрезан, полный объём в файле]"
                metadata["truncated"] = True
            
            metadata["chars_extracted"] = len(text)
            return text, metadata
            
        except ImportError as e:
            logger.warning(f"Отсутствует библиотека: {e}")
            return f"⚠️ Нужна библиотека: {e}. Установи через pip.", metadata
        except Exception as e:
            logger.error(f"Ошибка парсинга {file_path}: {e}")
            return f"❌ Ошибка парсинга: {e}", metadata
    
    @staticmethod
    async def _parse_pdf(file_path: str) -> str:
        """Извлечение текста из PDF через PyPDF2 или pdfplumber."""
        try:
            # Пробуем pdfplumber (лучше для таблиц)
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Страница {i+1} ---\n{page_text}")
                    
                    # Извлекаем таблицы если есть
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            # Формируем текстовую таблицу
                            rows = [" | ".join([str(cell or "") for cell in row]) for row in table]
                            text_parts.append("📊 Таблица:\n" + "\n".join(rows))
            
            return "\n\n".join(text_parts) if text_parts else "PDF пуст или содержит только изображения."
        
        except ImportError:
            # Fallback на PyPDF2
            try:
                import PyPDF2
                text_parts = []
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for i, page in enumerate(reader.pages):
                        text = page.extract_text()
                        if text:
                            text_parts.append(f"--- Страница {i+1} ---\n{text}")
                return "\n\n".join(text_parts) if text_parts else "PDF пуст."
            except ImportError:
                raise ImportError("pdfplumber или PyPDF2 (pip install pdfplumber)")
    
    @staticmethod
    async def _parse_docx(file_path: str) -> str:
        """Извлечение текста из DOCX через python-docx."""
        try:
            from docx import Document
            doc = Document(file_path)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Учитываем стили (заголовки)
                    if para.style and para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        text_parts.append(f"{'#' * int(level)} {para.text}")
                    else:
                        text_parts.append(para.text)
            
            # Извлекаем таблицы
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    text_parts.append("📊 Таблица:\n" + "\n".join(rows))
            
            return "\n".join(text_parts) if text_parts else "Документ пуст."
        except ImportError:
            raise ImportError("python-docx (pip install python-docx)")
    
    @staticmethod
    async def _parse_xlsx(file_path: str) -> str:
        """Извлечение данных из Excel через openpyxl."""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True)
            
            text_parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                text_parts.append(f"📄 Лист: {sheet_name}")
                
                rows_data = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(cell or "") for cell in row]
                    rows_data.append(" | ".join(cells))
                
                if rows_data:
                    # Ограничиваем до 100 строк на лист
                    if len(rows_data) > 100:
                        text_parts.append(f"(показаны первые 100 из {len(rows_data)} строк)")
                        rows_data = rows_data[:100]
                    text_parts.append("\n".join(rows_data))
            
            wb.close()
            return "\n\n".join(text_parts) if text_parts else "Таблица пуста."
        except ImportError:
            raise ImportError("openpyxl (pip install openpyxl)")
    
    @staticmethod
    async def _parse_csv(file_path: str) -> str:
        """Парсинг CSV файла."""
        import csv
        
        text_parts = []
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                if i > 200:  # Ограничение
                    text_parts.append(f"... (ещё {i} строк)")
                    break
                text_parts.append(" | ".join(row))
        
        return "\n".join(text_parts) if text_parts else "CSV пуст."
    
    @staticmethod
    async def _parse_text(file_path: str) -> str:
        """Чтение текстового файла как есть."""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
